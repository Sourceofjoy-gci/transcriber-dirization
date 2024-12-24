import streamlit as st
import whisper
import datetime
import torch
import numpy as np
from sklearn.cluster import AgglomerativeClustering
import wave
import contextlib
import tempfile
import os
from pathlib import Path
import subprocess
import torchaudio
from torch.nn import functional as F
import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from agents import (
    TranscriptSegment,
    TranscriptMetadata,
    SpeakerIdentificationAgent,
    FormattingStructureAgent,
    LanguageGrammarRefinementAgent,
    QualityAssuranceAgent
)
from live_recording import LiveAudioRecorder, get_available_devices

class AudioEmbedding:
    def __init__(self):
        """Initialize audio embedding with advanced processing"""
        self.sample_rate = 16000
        self.embedding_size = 192
        self.window_size = 0.5  # 500ms windows for analysis
        self.min_speech_duration = 0.3  # Minimum speech segment duration
        
    def process_waveform(self, waveform):
        """Enhanced audio preprocessing with noise reduction"""
        # Convert to mono if stereo
        if waveform.shape[0] > 1:
            waveform = torch.mean(waveform, dim=0, keepdim=True)
        
        # Resample to target rate
        if self.sample_rate != 16000:
            resampler = torchaudio.transforms.Resample(
                orig_freq=self.sample_rate, new_freq=16000
            )
            waveform = resampler(waveform)
        
        # Apply preprocessing
        waveform = self._apply_noise_reduction(waveform)
        waveform = self._normalize_audio(waveform)
        
        return waveform

    def get_embedding(self, waveform):
        """Generate advanced audio embeddings for speaker diarization"""
        waveform = self.process_waveform(waveform)
        
        with torch.no_grad():
            # Extract multiple features
            features = []
            
            # 1. Mel Spectrogram features
            mel_spec = torchaudio.transforms.MelSpectrogram(
                sample_rate=self.sample_rate,
                n_mels=80,
                n_fft=400,
                hop_length=160
            )(waveform)
            mel_features = torch.mean(mel_spec, dim=2)  # [batch, n_mels]
            features.append(mel_features)
            
            # 2. MFCC features
            mfcc_transform = torchaudio.transforms.MFCC(
                sample_rate=self.sample_rate,
                n_mfcc=40
            )
            mfcc = mfcc_transform(waveform)
            mfcc_features = torch.mean(mfcc, dim=2)  # [batch, n_mfcc]
            features.append(mfcc_features)
            
            # 3. Pitch features
            pitch = self._extract_pitch_features(waveform)  # [batch, 1, time]
            pitch_features = torch.mean(pitch, dim=2)  # [batch, 1]
            features.append(pitch_features)
            
            # Ensure all features have the same dimensions before concatenating
            features = [f.view(1, -1) if len(f.shape) == 1 else f for f in features]
            
            # Combine features
            combined_features = torch.cat(features, dim=1)  # [batch, total_features]
            normalized_features = F.normalize(combined_features, p=2, dim=1)
            
            # Ensure consistent embedding size
            if normalized_features.shape[1] > self.embedding_size:
                normalized_features = normalized_features[:, :self.embedding_size]
            else:
                padding_size = self.embedding_size - normalized_features.shape[1]
                normalized_features = F.pad(normalized_features, (0, padding_size))
            
        return normalized_features.numpy().reshape(-1)

    def _apply_noise_reduction(self, waveform):
        """Apply noise reduction to improve signal quality"""
        # Spectral subtraction-based noise reduction
        spec = torch.stft(
            waveform[0], 
            n_fft=400, 
            hop_length=160, 
            window=torch.hann_window(400).to(waveform.device),
            return_complex=True
        )
        
        # Estimate noise from non-speech segments
        mag_spec = torch.abs(spec)
        noise_estimate = torch.mean(mag_spec[:, :50], dim=1, keepdim=True)  # Use first 50 frames
        
        # Apply spectral subtraction
        mag_spec_clean = torch.maximum(mag_spec - 2 * noise_estimate, torch.zeros_like(mag_spec))
        phase_spec = torch.angle(spec)
        spec_clean = mag_spec_clean * torch.exp(1j * phase_spec)
        
        # Convert back to time domain
        waveform_clean = torch.istft(
            spec_clean,
            n_fft=400,
            hop_length=160,
            window=torch.hann_window(400).to(waveform.device)
        )
        
        return waveform_clean.unsqueeze(0)

    def _normalize_audio(self, waveform):
        """Normalize audio levels"""
        return waveform / (torch.max(torch.abs(waveform)) + 1e-8)

    def _extract_pitch_features(self, waveform):
        """Extract pitch-related features"""
        # Simple pitch detection using autocorrelation
        frame_length = int(self.window_size * self.sample_rate)
        hop_length = frame_length // 2
        
        # Ensure waveform is 2D [channels, samples]
        if len(waveform.shape) == 1:
            waveform = waveform.unsqueeze(0)
        elif len(waveform.shape) == 3:
            waveform = waveform.squeeze(1)
            
        # Create overlapping frames
        num_frames = (waveform.shape[1] - frame_length) // hop_length + 1
        frames = torch.zeros((num_frames, frame_length))
        
        for i in range(num_frames):
            start = i * hop_length
            frames[i] = waveform[0, start:start + frame_length]
        
        # Normalize frames
        frames = frames / (torch.max(torch.abs(frames), dim=1, keepdim=True)[0] + 1e-8)
        
        # Calculate autocorrelation using FFT
        fft_size = 2 ** (frame_length - 1).bit_length()  # Next power of 2
        fft = torch.fft.rfft(frames, n=fft_size)
        power_spectrum = torch.abs(fft) ** 2
        autocorr = torch.fft.irfft(power_spectrum)[:, :frame_length]
        
        # Find peaks in autocorrelation (exclude first peak at lag 0)
        start_idx = int(self.sample_rate / 500)  # Min frequency = 500Hz
        end_idx = int(self.sample_rate / 50)     # Max frequency = 50Hz
        peaks = torch.argmax(autocorr[:, start_idx:end_idx], dim=1) + start_idx
        
        # Convert peak positions to pitch values (in Hz)
        pitch = self.sample_rate / (peaks.float() + 1)
        
        # Reshape to match expected dimensions [batch, features, time]
        pitch = pitch.unsqueeze(0).unsqueeze(0)
        
        return pitch

def convert_to_wav(input_path, output_path):
    """Convert uploaded audio file to WAV format."""
    try:
        subprocess.call(['ffmpeg', '-i', input_path, '-acodec', 'pcm_s16le', '-ar', '16000', '-ac', '1', output_path, '-y'])
        return True
    except Exception as e:
        st.error(f"Error converting audio: {str(e)}")
        return False

def get_audio_duration(path):
    """Get duration of audio file in seconds."""
    with contextlib.closing(wave.open(path, 'r')) as f:
        frames = f.getnframes()
        rate = f.getframerate()
        return frames / float(rate)

def load_audio_segment(path, start_time, end_time):
    """Load audio segment using torchaudio and convert to bytes for playback."""
    try:
        import torchaudio
        import io
        import soundfile as sf
        import numpy as np

        # Load the full audio file
        waveform, sample_rate = torchaudio.load(path)
        
        # Calculate frame indices
        start_frame = int(start_time * sample_rate)
        end_frame = int(end_time * sample_rate)
        
        # Extract segment
        segment = waveform[:, start_frame:end_frame]
        
        # Convert to numpy and then to bytes
        audio_numpy = segment.numpy()
        
        # If stereo, convert to mono by averaging channels
        if audio_numpy.shape[0] > 1:
            audio_numpy = np.mean(audio_numpy, axis=0)
        else:
            audio_numpy = audio_numpy[0]

        # Create bytes buffer
        buffer = io.BytesIO()
        
        # Write to buffer in WAV format
        sf.write(buffer, audio_numpy, sample_rate, format='WAV')
        
        # Get the bytes
        buffer.seek(0)
        return buffer

    except Exception as e:
        st.error(f"Error loading audio segment: {str(e)}")
        return None

def process_audio(audio_path, num_speakers, model_size='base'):
    """Process audio file for transcription and speaker diarization."""
    try:
        # Initialize components
        audio_embedder = AudioEmbedding()
        sia = SpeakerIdentificationAgent()
        fsa = FormattingStructureAgent()
        lgra = LanguageGrammarRefinementAgent()
        qaa = QualityAssuranceAgent()
        
        # Load and process audio
        waveform, sample_rate = torchaudio.load(audio_path)
        
        # Generate embeddings for speaker diarization
        embeddings = []
        window_size = int(audio_embedder.window_size * sample_rate)
        hop_length = window_size // 2
        
        for i in range(0, waveform.shape[1] - window_size, hop_length):
            segment = waveform[:, i:i + window_size]
            embedding = audio_embedder.get_embedding(segment)
            embeddings.append(embedding)
        
        embeddings = np.array(embeddings)
        
        # Normalize embeddings for better clustering
        embeddings_norm = np.linalg.norm(embeddings, axis=1, keepdims=True)
        embeddings_normalized = embeddings / (embeddings_norm + 1e-8)
        
        # Perform speaker clustering with compatible parameters
        clustering = AgglomerativeClustering(
            n_clusters=num_speakers,
            metric='euclidean',
            linkage='average'  # Using average linkage which works better with normalized embeddings
        )
        labels = clustering.fit_predict(embeddings_normalized)
        
        # Apply median filtering to smooth labels
        smoothed_labels = _smooth_speaker_labels(labels)
        
        # Load whisper model for transcription
        model = whisper.load_model(model_size)
        
        # Transcribe audio
        result = model.transcribe(audio_path)
        
        # Create transcript segments with improved speaker assignment
        segments = []
        for i, segment in enumerate(result["segments"]):
            # Find corresponding speaker label
            start_frame = int(segment["start"] / audio_embedder.window_size)
            end_frame = int(segment["end"] / audio_embedder.window_size)
            if start_frame >= len(smoothed_labels):
                start_frame = len(smoothed_labels) - 1
            if end_frame >= len(smoothed_labels):
                end_frame = len(smoothed_labels) - 1
                
            speaker_label = _get_dominant_speaker(smoothed_labels[start_frame:end_frame + 1])
            
            original_text = segment["text"]
            refined_text = lgra.refine_text(original_text)
            
            segments.append(TranscriptSegment(
                speaker=f"Speaker {speaker_label}",
                text=refined_text,
                start_time=segment["start"],
                end_time=segment["end"],
                confidence=segment.get("confidence", 1.0),
                original_text=original_text
            ))
        
        # Process segments through speaker identification
        processed_segments = sia.identify_speakers(segments)
        
        # Extract speaker mapping from the agent
        speaker_mapping = sia.speaker_mapping
        
        # Format transcript with processed segments
        formatted_transcript = fsa.format_transcript(processed_segments, speaker_mapping)
        
        # Quality assurance check
        qa_report = qaa.validate_transcript(processed_segments, speaker_mapping)
        
        return {
            "transcript": formatted_transcript,
            "speaker_mapping": speaker_mapping,
            "qa_report": qa_report,
            "segments": processed_segments,
            "audio_path": audio_path
        }
        
    except Exception as e:
        st.error(f"Error processing audio: {str(e)}")
        return None

def create_docx_transcript(segments, speaker_mapping, qa_report):
    """Create a formatted DOCX document from the transcript segments."""
    doc = Document()
    
    # Add title
    title = doc.add_heading('OFFICIAL COURT TRANSCRIPT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    doc.add_paragraph(f"Date: {datetime.datetime.now().strftime('%B %d, %Y')}")
    doc.add_paragraph(f"Time: {datetime.datetime.now().strftime('%H:%M:%S')}")
    doc.add_paragraph("")
    
    # Add quality report summary
    doc.add_heading('Quality Assurance Summary', level=1)
    for summary_item in qa_report['report_summary']:
        doc.add_paragraph(summary_item, style='List Bullet')
    doc.add_paragraph("")
    
    # Add speaker mapping
    doc.add_heading('Speaker Identification', level=1)
    for speaker, role in speaker_mapping.items():
        doc.add_paragraph(f"{speaker} → {role}", style='List Bullet')
    doc.add_paragraph("")
    
    # Add transcript content
    doc.add_heading('Transcript', level=1)
    for segment in segments:
        speaker_role = speaker_mapping.get(segment.speaker, segment.speaker)
        timestamp = f"[{segment.start_time:.2f}s - {segment.end_time:.2f}s]"
        
        p = doc.add_paragraph()
        speaker_run = p.add_run(f"{speaker_role}: ")
        speaker_run.bold = True
        p.add_run(f"{timestamp}\n{segment.text}")
        doc.add_paragraph("")
    
    return doc

def main():
    st.title("Court Transcript Analysis System")
    
    # Add input method selection
    input_method = st.radio(
        "Choose input method",
        ["Upload Audio File", "Live Recording"]
    )
    
    if input_method == "Upload Audio File":
        st.write("Upload an audio file for automated transcription and analysis")
        uploaded_file = st.file_uploader("Choose an audio file", type=['wav', 'mp3', 'm4a'])
        st.caption("Maximum file size: 700MB")
        
        if uploaded_file is not None:
            with st.spinner("Processing audio file..."):
                # Save uploaded file temporarily
                with tempfile.NamedTemporaryFile(delete=False, suffix=Path(uploaded_file.name).suffix) as tmp_file:
                    tmp_file.write(uploaded_file.getvalue())
                    audio_path = tmp_file.name
                
                # Convert to WAV if needed
                if not audio_path.endswith('.wav'):
                    wav_path = audio_path + '.wav'
                    convert_to_wav(audio_path, wav_path)
                    audio_path = wav_path
                
                # Process audio
                num_speakers = st.number_input("Estimated number of speakers", min_value=1, max_value=10, value=2)
                model_size = st.selectbox("Model size", ['tiny', 'base', 'small', 'medium', 'large'])
                
                if st.button("Process Audio"):
                    result = process_audio(audio_path, num_speakers, model_size)
                    
                    if result:
                        # Display transcript
                        st.subheader("Transcript")
                        for segment in result["segments"]:
                            speaker_role = result["speaker_mapping"].get(segment.speaker, segment.speaker)
                            st.markdown(f"**[{_format_timestamp(segment.start_time)}] {speaker_role}:**")
                            st.write(segment.original_text)
                            st.markdown("---")
                        
                        # Create and offer download of DOCX
                        doc = create_docx_transcript(
                            result["segments"],
                            result["speaker_mapping"],
                            result["qa_report"]
                        )
                        
                        if doc:
                            # Convert to bytes and offer download
                            docx_bytes = io.BytesIO()
                            doc.save(docx_bytes)
                            docx_bytes = docx_bytes.getvalue()
                            
                            st.download_button(
                                label="Download Transcript (DOCX)",
                                data=docx_bytes,
                                file_name="court_transcript.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                            
                            # Create and offer download of TXT
                            txt_content = ""
                            for segment in result["segments"]:
                                speaker_role = result["speaker_mapping"].get(segment.speaker, segment.speaker)
                                txt_content += f"[{_format_timestamp(segment.start_time)}] {speaker_role}: {segment.original_text}\n\n"
                            
                            st.download_button(
                                label="Download Transcript (TXT)",
                                data=txt_content,
                                file_name="court_transcript.txt",
                                mime="text/plain"
                            )
                            
                            st.success("Transcription complete! You can now download the transcript in DOCX or TXT format.")
            
            # Cleanup temporary files
            try:
                os.remove(audio_path)
                if audio_path.endswith('.wav') and not uploaded_file.name.endswith('.wav'):
                    os.remove(audio_path[:-4])  # Remove original temp file
            except:
                pass
    
    else:  # Live Recording
        st.write("Record audio live for transcription and analysis")
        
        if 'recorder' not in st.session_state:
            st.session_state.recorder = None
            st.session_state.recording = False
            st.session_state.audio_path = None
        
        # Get available audio devices
        available_devices = get_available_devices()
        device_options = {f"{device['name']} (Sample Rate: {device['sample_rate']}Hz)": device['index'] 
                         for device in available_devices}
        
        # Device selection
        selected_device_name = st.selectbox(
            "Select Audio Input Device",
            options=list(device_options.keys()),
            index=0
        )
        selected_device_index = device_options[selected_device_name]
        
        col1, col2 = st.columns(2)
        
        with col1:
            if not st.session_state.recording:
                if st.button("Start Recording"):
                    st.session_state.recorder = LiveAudioRecorder(device_index=selected_device_index)
                    st.session_state.recorder.start_recording()
                    st.session_state.recording = True
                    st.rerun()
            
        with col2:
            if st.session_state.recording:
                if st.button("Stop Recording"):
                    audio_path = st.session_state.recorder.stop_recording()
                    st.session_state.recorder.close()
                    st.session_state.recording = False
                    st.session_state.audio_path = audio_path
                    st.rerun()
        
        if st.session_state.recording:
            st.warning("Recording in progress...")
        
        if st.session_state.audio_path:
            st.success(f"Recording saved as: {st.session_state.audio_path}")
            
            # Process the recorded audio
            num_speakers = st.number_input("Estimated number of speakers", min_value=1, max_value=10, value=2)
            model_size = st.selectbox("Model size", ['tiny', 'base', 'small', 'medium', 'large'])
            
            if st.button("Process Recording"):
                result = process_audio(st.session_state.audio_path, num_speakers, model_size)
                
                if result:
                    # Display transcript
                    st.subheader("Transcript")
                    for segment in result["segments"]:
                        speaker_role = result["speaker_mapping"].get(segment.speaker, segment.speaker)
                        st.markdown(f"**[{_format_timestamp(segment.start_time)}] {speaker_role}:**")
                        st.write(segment.original_text)
                        st.markdown("---")
                    
                    # Create and offer download of DOCX
                    doc = create_docx_transcript(
                        result["segments"],
                        result["speaker_mapping"],
                        result["qa_report"]
                    )
                    
                    if doc:
                        # Convert to bytes and offer download
                        docx_bytes = io.BytesIO()
                        doc.save(docx_bytes)
                        docx_bytes = docx_bytes.getvalue()
                        
                        st.download_button(
                            label="Download Transcript (DOCX)",
                            data=docx_bytes,
                            file_name="court_transcript.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                        
                        # Create and offer download of TXT
                        txt_content = ""
                        for segment in result["segments"]:
                            speaker_role = result["speaker_mapping"].get(segment.speaker, segment.speaker)
                            txt_content += f"[{_format_timestamp(segment.start_time)}] {speaker_role}: {segment.original_text}\n\n"
                        
                        st.download_button(
                            label="Download Transcript (TXT)",
                            data=txt_content,
                            file_name="court_transcript.txt",
                            mime="text/plain"
                        )
                        
                        st.success("Transcription complete! You can now download the transcript in DOCX or TXT format.")

def _format_timestamp(seconds: float) -> str:
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    seconds = int(seconds % 60)
    return f"{hours:02d}:{minutes:02d}:{seconds:02d}"

def _smooth_speaker_labels(labels):
    # Simple smoothing by majority vote
    smoothed_labels = []
    window_size = 10
    for i in range(len(labels)):
        window = labels[max(0, i - window_size):i + window_size + 1]
        smoothed_labels.append(np.bincount(window).argmax())
    return smoothed_labels

def _get_dominant_speaker(labels):
    # Get the most frequent speaker label
    return np.bincount(labels).argmax()

if __name__ == "__main__":
    main()